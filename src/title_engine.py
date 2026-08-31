# Файл: src/title_engine.py
"""
Генератор титульных листов с жестким соблюдением сетки ГОСТ и защитой от съезжания верстки.
"""

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from src.schema import DocProfile, TitleDefaults


def create_title_page(doc: Document, profile: DocProfile, data: TitleDefaults):
    """
    Генерирует первую страницу (титульный лист) в первом разделе документа.
    """
    # Настраиваем первый раздел (титульный)
    sec_title = doc.sections[0]
    sec_title.page_width = Mm(210)
    sec_title.page_height = Mm(297)
    sec_title.left_margin = Mm(profile.page.margins.left_mm)
    sec_title.right_margin = Mm(profile.page.margins.right_mm)
    sec_title.top_margin = Mm(profile.page.margins.top_mm)
    sec_title.bottom_margin = Mm(profile.page.margins.bottom_mm)

    # Отключаем колонтитулы на титульнике
    sec_title.different_first_page_header_footer = True
    sec_title.header.is_linked_to_previous = False
    sec_title.footer.is_linked_to_previous = False
    for p in sec_title.header.paragraphs:
        p.text = ""
    for p in sec_title.footer.paragraphs:
        p.text = ""

    # 1. Шапка: Министерство
    p_min = doc.add_paragraph()
    p_min.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_min.paragraph_format.space_before = Pt(0)
    p_min.paragraph_format.space_after = Pt(2)
    p_min.paragraph_format.line_spacing = 1.15
    run_min = p_min.add_run(data.ministry)
    run_min.font.name = profile.body.font_family
    run_min.font.size = Pt(11)
    run_min.font.bold = True

    # 2. Учебное заведение
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_before = Pt(2)
    p_inst.paragraph_format.space_after = Pt(4)
    p_inst.paragraph_format.line_spacing = 1.15
    run_inst = p_inst.add_run(data.institution)
    run_inst.font.name = profile.body.font_family
    run_inst.font.size = Pt(13)
    run_inst.font.bold = True

    # 3. Отделение / Факультет
    if data.department:
        p_dep = doc.add_paragraph()
        p_dep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_dep.paragraph_format.space_before = Pt(0)
        p_dep.paragraph_format.space_after = Pt(2)
        run_dep = p_dep.add_run(data.department)
        run_dep.font.name = profile.body.font_family
        run_dep.font.size = Pt(12)

    # 4. Специальность
    p_spec = doc.add_paragraph()
    p_spec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_spec.paragraph_format.space_before = Pt(2)
    p_spec.paragraph_format.space_after = Pt(36)
    run_spec = p_spec.add_run(f"Специальность: {data.specialty}")
    run_spec.font.name = profile.body.font_family
    run_spec.font.size = Pt(12)

    # 5. Тип работы (КУРСОВОЙ ПРОЕКТ / ДИПЛОМНЫЙ ПРОЕКТ)
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_type.paragraph_format.space_before = Pt(24)
    p_type.paragraph_format.space_after = Pt(6)
    run_type = p_type.add_run(data.work_type.upper())
    run_type.font.name = profile.body.font_family
    run_type.font.size = Pt(16)
    run_type.font.bold = True

    # 6. Дисциплина (если указана)
    if data.discipline:
        p_disc = doc.add_paragraph()
        p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_disc.paragraph_format.space_before = Pt(0)
        p_disc.paragraph_format.space_after = Pt(12)
        run_disc = p_disc.add_run(f"по МДК / дисциплине: «{data.discipline}»")
        run_disc.font.name = profile.body.font_family
        run_disc.font.size = Pt(13)

    # 7. Тема проекта
    p_theme = doc.add_paragraph()
    p_theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_theme.paragraph_format.space_before = Pt(12)
    p_theme.paragraph_format.space_after = Pt(36)
    run_theme_label = p_theme.add_run("на тему: ")
    run_theme_label.font.name = profile.body.font_family
    run_theme_label.font.size = Pt(14)
    run_theme = p_theme.add_run(f"«{data.theme}»")
    run_theme.font.name = profile.body.font_family
    run_theme.font.size = Pt(14)
    run_theme.font.bold = True

    # 8. Блок подписей (Таблица из 2 колонок без границ, прижатая вправо)
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    tbl.autofit = False

    # Задаем ширину колонок
    col_widths = [Mm(40), Mm(80)]
    for row in tbl.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Убираем все границы у таблицы
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders %s>
                    <w:top w:val="none"/>
                    <w:left w:val="none"/>
                    <w:bottom w:val="none"/>
                    <w:right w:val="none"/>
                </w:tcBorders>
            ''' % nsdecls('w'))
            tcPr.append(tcBorders)

    # Заполняем "Выполнил"
    cell_00 = tbl.cell(0, 0)
    p_00 = cell_00.paragraphs[0]
    p_00.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_00.add_run("Выполнил:")
    r.font.name = profile.body.font_family
    r.font.size = Pt(12)

    cell_01 = tbl.cell(0, 1)
    p_01 = cell_01.paragraphs[0]
    p_01.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_01.add_run(f"студент {data.student_group}\n{data.student_name}")
    r.font.name = profile.body.font_family
    r.font.size = Pt(12)
    r.font.bold = True

    # Заполняем "Проверил"
    cell_10 = tbl.cell(1, 0)
    p_10 = cell_10.paragraphs[0]
    p_10.paragraph_format.space_before = Pt(8)
    p_10.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_10.add_run("Проверил:")
    r.font.name = profile.body.font_family
    r.font.size = Pt(12)

    cell_11 = tbl.cell(1, 1)
    p_11 = cell_11.paragraphs[0]
    p_11.paragraph_format.space_before = Pt(8)
    p_11.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_11.add_run(f"{data.supervisor_title}\n{data.supervisor_name}")
    r.font.name = profile.body.font_family
    r.font.size = Pt(12)
    r.font.bold = True

    # 9. Город и год внизу
    p_bottom = doc.add_paragraph()
    p_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bottom.paragraph_format.space_before = Pt(48)
    p_bottom.paragraph_format.space_after = Pt(0)
    r_bottom = p_bottom.add_run(f"{data.city} – {data.year}")
    r_bottom.font.name = profile.body.font_family
    r_bottom.font.size = Pt(12)

    # 10. Добавляем новый раздел (Section 2) для основного контента
    sec_main = doc.add_section()
    sec_main.page_width = Mm(210)
    sec_main.page_height = Mm(297)
    sec_main.left_margin = Mm(profile.page.margins.left_mm)
    sec_main.right_margin = Mm(profile.page.margins.right_mm)
    sec_main.top_margin = Mm(profile.page.margins.top_mm)
    sec_main.bottom_margin = Mm(profile.page.margins.bottom_mm)

    # Настраиваем нумерацию во 2 разделе
    if profile.page.page_numbering.enabled:
        setup_page_numbering(sec_main, profile)

    return sec_main


def setup_page_numbering(section, profile: DocProfile):
    """
    Добавляет нативное поле нумерации страницы (w:fldSimple w:instr="PAGE") в правый верхний угол.
    """
    section.header.is_linked_to_previous = False
    p_head = section.header.paragraphs[0]
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_head.paragraph_format.space_before = Pt(0)
    p_head.paragraph_format.space_after = Pt(0)

    # Вставляем XML поле страницы PAGE
    r = p_head.add_run()
    r.font.name = profile.page.page_numbering.font_family
    r.font.size = Pt(profile.page.page_numbering.font_size_pt)
    
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    r._r.append(fldSimple)
