# Файл: /home/detker/Документы/repository/doc-forge/src/toc_engine.py
"""
Генератор оглавления (СОДЕРЖАНИЕ) по ГОСТу с отточиями (leader dots) и точными номерами страниц.
"""

from typing import List, Tuple
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from src.schema import DocProfile


def create_table_of_contents(
    doc: Document,
    profile: DocProfile,
    items: List[Tuple[str, int, int]]  # (title, level, page_num)
):
    """
    Генерирует страницу «СОДЕРЖАНИЕ» с отточиями и выравниванием номеров страниц по правому краю.
    """
    # Заголовок СОДЕРЖАНИЕ
    p_toc_head = doc.add_paragraph()
    p_toc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_head.paragraph_format.space_before = Pt(0)
    p_toc_head.paragraph_format.space_after = Pt(profile.headings.h1.spacing_after_pt)
    p_toc_head.paragraph_format.line_spacing = 1.5
    
    r_head = p_toc_head.add_run("СОДЕРЖАНИЕ")
    r_head.font.name = profile.body.font_family
    r_head.font.size = Pt(profile.headings.h1.font_size_pt)
    r_head.font.bold = True

    # Вычисляем позицию правого края для табуляции (210мм - левое поле - правое поле)
    content_width_mm = 210.0 - profile.page.margins.left_mm - profile.page.margins.right_mm
    tab_pos = Mm(content_width_mm)

    for title, level, page_num in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Добавляем Tab Stop с точками (отточиями) на правом краю
        p.paragraph_format.tab_stops.add_tab_stop(
            tab_pos,
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS
        )

        # Отступ для подразделов (H2)
        if level == 2:
            p.paragraph_format.left_indent = Mm(profile.headings.h2.first_line_indent_cm * 10)
        elif level == 3:
            p.paragraph_format.left_indent = Mm((profile.headings.h3.first_line_indent_cm + 0.5) * 10)

        # Текст элемента
        run_title = p.add_run(title)
        run_title.font.name = profile.body.font_family
        run_title.font.size = Pt(profile.body.font_size_pt)
        if level == 1:
            run_title.font.bold = True

        # Символ табуляции (вызовет отрисовку точек до правого поля)
        run_tab = p.add_run("\t")
        run_tab.font.name = profile.body.font_family
        run_tab.font.size = Pt(profile.body.font_size_pt)

        # Номер страницы
        run_page = p.add_run(str(page_num))
        run_page.font.name = profile.body.font_family
        run_page.font.size = Pt(profile.body.font_size_pt)
        if level == 1:
            run_page.font.bold = True

    # Разрыв страницы после оглавления
    doc.add_page_break()
