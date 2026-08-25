# Файл: /home/detker/Документы/repository/doc-forge/src/compiler.py
"""
Компилятор Markdown -> DOCX с применением профиля оформления ГОСТ.
"""

import re
import yaml
from pathlib import Path
from typing import Dict, Any, Tuple, Optional, List

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from src.schema import DocProfile, TitleDefaults, TextStyle
from src.title_engine import create_title_page
from src.toc_engine import create_table_of_contents


def apply_paragraph_style(p, style: TextStyle, doc_profile: DocProfile, is_first_heading: bool = False):
    """
    Применяет точные метрики шрифта и отступов к параграфу docx.
    """
    pf = p.paragraph_format
    pf.line_spacing = style.line_spacing
    pf.space_before = Pt(style.spacing_before_pt)
    pf.space_after = Pt(style.spacing_after_pt)
    
    if style.first_line_indent_cm > 0:
        pf.first_line_indent = Mm(style.first_line_indent_cm * 10)
    else:
        pf.first_line_indent = Mm(0)

    # Выравнивание
    align_map = {
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    p.alignment = align_map.get(style.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Запрет висячих строк и отрыва заголовков
    if style.keep_with_next:
        pf.keep_with_next = True
    pf.widow_control = True


def parse_frontmatter(md_text: str) -> Tuple[Dict[str, Any], str]:
    """
    Извлекает YAML Frontmatter из начала Markdown файла.
    """
    # Удаляем служебные комментарии в самом начале файла
    md_text = re.sub(r'^(#\s*Файл:.*?\n|<!--.*?-->\n)+', '', md_text.strip(), flags=re.DOTALL)
    
    if md_text.startswith("---"):
        parts = md_text.split("---", 2)
        if len(parts) >= 3:
            try:
                fm_data = yaml.safe_load(parts[1]) or {}
                return fm_data, parts[2].strip()
            except Exception:
                pass
    return {}, md_text.strip()


def add_formatted_runs(p, text: str, base_style: TextStyle, all_caps: bool = False):
    """
    Парсит жирный текст **жирный** и курсив *курсив* внутри строки.
    """
    if all_caps:
        text = text.upper()

    # Токенизатор inline Markdown (**bold**, *italic*)
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|[^\*`]+)')
    tokens = pattern.findall(text)

    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) >= 4:
            clean_text = token[2:-2]
            r = p.add_run(clean_text)
            r.font.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) >= 2:
            clean_text = token[1:-1]
            r = p.add_run(clean_text)
            r.font.italic = True
        elif token.startswith("`") and token.endswith("`") and len(token) >= 2:
            clean_text = token[1:-1]
            r = p.add_run(clean_text)
            r.font.name = "Courier New"
            r.font.size = Pt(base_style.font_size_pt - 1)
            continue
        else:
            r = p.add_run(token)
            if base_style.bold:
                r.font.bold = True
            if base_style.italic:
                r.font.italic = True

        r.font.name = base_style.font_family
        r.font.size = Pt(base_style.font_size_pt)


def set_cell_margins_and_borders(cell, top=100, bottom=100, left=150, right=150):
    """
    Устанавливает внутренние отступы (padding) и тонкие черные границы для ячейки таблицы.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders %s>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tcBorders>
    ''' % nsdecls('w'))
    tcPr.append(tcBorders)

    tcMar = parse_xml(r'''
        <w:tcMar %s>
            <w:top w:w="%d" w:type="dxa"/>
            <w:bottom w:w="%d" w:type="dxa"/>
            <w:left w:w="%d" w:type="dxa"/>
            <w:right w:w="%d" w:type="dxa"/>
        </w:tcMar>
    ''' % (nsdecls('w'), top, bottom, left, right))
    tcPr.append(tcMar)


def build_document(md_content: str, profile: DocProfile, output_path: Path):
    """
    Главная функция сборки Markdown -> DOCX с применением профиля.
    """
    fm, body_md = parse_frontmatter(md_content)

    # Обновляем метаданные титульного листа из Frontmatter, если есть
    title_data = profile.title_defaults.model_copy()
    if "title" in fm and isinstance(fm["title"], dict):
        for k, v in fm["title"].items():
            if hasattr(title_data, k):
                setattr(title_data, k, v)

    doc = Document()

    # 1. Генерируем титульный лист (Стр. 1)
    create_title_page(doc, profile, title_data)

    # 2. Генерируем оглавление (Стр. 2), если есть
    toc_items = []
    if "toc" in fm and isinstance(fm["toc"], list):
        for item in fm["toc"]:
            toc_items.append((item.get("title", ""), item.get("level", 1), item.get("page", 3)))
        create_table_of_contents(doc, profile, toc_items)

    # 3. Парсим строки Markdown
    lines = body_md.split("\n")
    i = 0
    table_caption = ""
    first_h1 = True

    while i < len(lines):
        line = lines[i].strip()
        
        # Пропуск пустых строк (защита от лишних w:p)
        if not line or line.startswith("<!--") or line.startswith("# Файл:"):
            i += 1
            continue

        # Разрыв страницы
        if line == "---" or line == "<pagebreak>":
            doc.add_page_break()
            i += 1
            continue

        # Заголовок 1 уровня (# ЗАГОЛОВОК)
        if line.startswith("# ") and not line.startswith("## "):
            heading_text = line[2:].strip()
            p = doc.add_paragraph()
            style = profile.headings.h1
            # Не добавляем лишний разрыв страницы, если заголовок идет первым сразу после TOC
            apply_paragraph_style(p, style, profile, is_first_heading=first_h1)
            add_formatted_runs(p, heading_text, style, all_caps=style.all_caps)
            first_h1 = False
            i += 1
            continue

        # Заголовок 2 уровня (## 1.1 Заголовок)
        if line.startswith("## ") and not line.startswith("### "):
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            style = profile.headings.h2
            apply_paragraph_style(p, style, profile)
            add_formatted_runs(p, heading_text, style, all_caps=style.all_caps)
            i += 1
            continue

        # Заголовок 3 уровня (### 1.1.1 Пункт)
        if line.startswith("### "):
            heading_text = line[4:].strip()
            p = doc.add_paragraph()
            style = profile.headings.h3
            apply_paragraph_style(p, style, profile)
            add_formatted_runs(p, heading_text, style, all_caps=style.all_caps)
            i += 1
            continue

        # Подпись таблицы (например: "Таблица 1 — Название")
        if line.startswith("Таблица ") or line.startswith("Table "):
            table_caption = line
            i += 1
            continue

        # Начало таблицы Markdown (| a | b |)
        if line.startswith("|") and line.endswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                row_raw = lines[i].strip()
                # Пропускаем разделитель |---|---|
                if not re.match(r'^\|[\s\:\-\|]+\|$', row_raw):
                    cells = [c.strip() for c in row_raw[1:-1].split("|")]
                    table_rows.append(cells)
                i += 1

            if table_rows:
                # Вставляем подпись таблицы перед таблицей
                if table_caption:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_cap.paragraph_format.first_line_indent = Mm(profile.body.first_line_indent_cm * 10)
                    p_cap.paragraph_format.space_before = Pt(profile.elements.tables.spacing_before_pt)
                    p_cap.paragraph_format.space_after = Pt(2)
                    p_cap.paragraph_format.line_spacing = 1.15
                    r_cap = p_cap.add_run(table_caption)
                    r_cap.font.name = profile.elements.tables.font_family
                    r_cap.font.size = Pt(profile.elements.tables.font_size_pt)
                    table_caption = ""

                # Создаем саму таблицу
                num_rows = len(table_rows)
                num_cols = len(table_rows[0])
                tbl = doc.add_table(rows=num_rows, cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = True

                # Задаем повтор шапки на каждой новой странице (w:tblHeader)
                header_tr = tbl.rows[0]._tr.get_or_add_trPr()
                header_tr.append(parse_xml(r'<w:tblHeader %s/>' % nsdecls('w')))
                header_tr.append(parse_xml(r'<w:cantSplit %s/>' % nsdecls('w')))

                for r_idx, row_data in enumerate(table_rows):
                    trPr = tbl.rows[r_idx]._tr.get_or_add_trPr()
                    trPr.append(parse_xml(r'<w:cantSplit %s/>' % nsdecls('w')))
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx < num_cols:
                            cell = tbl.cell(r_idx, c_idx)
                            set_cell_margins_and_borders(cell)
                            p_cell = cell.paragraphs[0]
                            p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            p_cell.paragraph_format.space_before = Pt(2)
                            p_cell.paragraph_format.space_after = Pt(2)
                            p_cell.paragraph_format.line_spacing = profile.elements.tables.line_spacing
                            p_cell.paragraph_format.first_line_indent = Mm(0)
                            
                            r_c = p_cell.add_run(cell_value)
                            r_c.font.name = profile.elements.tables.font_family
                            r_c.font.size = Pt(profile.elements.tables.font_size_pt)
                            if r_idx == 0:
                                r_c.font.bold = True
                                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Отступ после таблицы
                p_after_tbl = doc.add_paragraph()
                p_after_tbl.paragraph_format.space_before = Pt(0)
                p_after_tbl.paragraph_format.space_after = Pt(profile.elements.tables.spacing_after_pt)
                p_after_tbl.paragraph_format.line_spacing = 1.0
                p_after_tbl.paragraph_format.first_line_indent = Mm(0)
            continue

        # Маркированные списки (- пункт или * пункт)
        if line.startswith("- ") or line.startswith("* "):
            item_text = line[2:].strip()
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = profile.body.line_spacing
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Mm(profile.elements.lists.first_line_indent_cm * 10)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Дефис по ГОСТу
            marker = "– " if profile.elements.lists.bullet_style == "dash" else "• "
            r_mark = p.add_run(marker)
            r_mark.font.name = profile.body.font_family
            r_mark.font.size = Pt(profile.body.font_size_pt)

            add_formatted_runs(p, item_text, profile.body)
            i += 1
            continue

        # Нумерованные списки (1. пункт)
        num_match = re.match(r'^(\d+[\.\)])\s+(.+)$', line)
        if num_match:
            num_prefix = num_match.group(1) + " "
            item_text = num_match.group(2).strip()
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = profile.body.line_spacing
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Mm(profile.elements.lists.first_line_indent_cm * 10)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            r_num = p.add_run(num_prefix)
            r_num.font.name = profile.body.font_family
            r_num.font.size = Pt(profile.body.font_size_pt)

            add_formatted_runs(p, item_text, profile.body)
            i += 1
            continue

        # Обычный параграф текста
        p = doc.add_paragraph()
        apply_paragraph_style(p, profile.body, profile)
        add_formatted_runs(p, line, profile.body)
        i += 1

    # Сохраняем файл
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
