# Файл: src/converters/docx_engine.py
"""
Движок глубокого анализа и конвертации файлов DOCX в семантический Markdown + профиль ГОСТ.
"""

import re
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from collections import Counter
import yaml

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

from src.converters.base import BaseConverter, ConversionResult
from src.schema import (
    DocProfile,
    MetaInfo,
    PageConfig,
    PageMargins,
    PageNumbering,
    TextStyle,
    HeadingStyles,
    ElementStyles,
    TableStyle,
    ListStyle,
    ReferenceStyle,
    TitleDefaults,
)


class DocxConverter(BaseConverter):
    """
    Конвертер DOCX -> Markdown с извлечением семантики титульных листов и точных метрик оформления.
    """

    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".docx", ".docm"]

    def convert(self, file_path: Path, extract_profile: bool = True) -> ConversionResult:
        if not file_path.exists():
            raise FileNotFoundError(f"Файл '{file_path}' не найден!")

        doc = Document(str(file_path))
        
        # 1. Извлекаем параметры страницы и геометрии
        margins_dict, page_config = self._extract_page_config(doc)

        # 2. Анализируем титульный лист (если есть)
        title_meta, title_paras_count = self._extract_title_page_metadata(doc)

        # 3. Анализируем шрифтовую и абзацную типографику
        body_style, headings_styles, table_style, list_style = self._analyze_typography(doc, title_paras_count)

        # 4. Собираем итоговый профиль
        profile_name = file_path.stem.lower().replace(" ", "_").replace("-", "_")
        institution_str = title_meta.get("institution") or "Учебное заведение"
        profile = DocProfile(
            meta=MetaInfo(
                name=profile_name,
                institution=institution_str,
                standard="ГОСТ / СТО (извлечено из DOCX)",
                description=f"Автоматически извлеченный профиль из файла {file_path.name}",
            ),
            page=page_config,
            body=body_style,
            headings=headings_styles,
            elements=ElementStyles(
                tables=table_style,
                lists=list_style,
                references=ReferenceStyle(
                    font_family=body_style.font_family,
                    font_size_pt=body_style.font_size_pt,
                    line_spacing=body_style.line_spacing,
                    first_line_indent_cm=body_style.first_line_indent_cm,
                ),
            ),
            title_defaults=TitleDefaults(**title_meta) if title_meta else TitleDefaults(),
        )

        # 5. Генерируем семантический Markdown
        md_text, headings_hierarchy, tables_count = self._convert_document_to_markdown(
            doc, title_meta, profile, title_paras_count
        )

        # 6. Формируем подробный отчет о макете для AI и инженера
        layout_report = self._generate_layout_report(profile, title_meta, tables_count, len(headings_hierarchy))

        return ConversionResult(
            markdown=md_text,
            profile=profile if extract_profile else None,
            metadata={
                "source_file": str(file_path),
                "title_metadata": title_meta,
                "tables_count": tables_count,
                "headings_count": len(headings_hierarchy),
                "format": "DOCX",
            },
            layout_report=layout_report,
            tables_count=tables_count,
            headings_hierarchy=headings_hierarchy,
        )

    def _extract_page_config(self, doc: Document) -> Tuple[Dict[str, float], PageConfig]:
        sec = doc.sections[0]
        left_mm = round(sec.left_margin.mm, 1) if sec.left_margin else 30.0
        right_mm = round(sec.right_margin.mm, 1) if sec.right_margin else 10.0
        top_mm = round(sec.top_margin.mm, 1) if sec.top_margin else 20.0
        bottom_mm = round(sec.bottom_margin.mm, 1) if sec.bottom_margin else 20.0

        margins_dict = {
            "left_mm": left_mm,
            "right_mm": right_mm,
            "top_mm": top_mm,
            "bottom_mm": bottom_mm,
        }

        # Определяем ориентацию
        orientation = "portrait"
        if sec.page_width and sec.page_height and sec.page_width > sec.page_height:
            orientation = "landscape"

        page_config = PageConfig(
            format="A4",
            orientation=orientation,
            margins=PageMargins(**margins_dict),
            page_numbering=PageNumbering(
                enabled=True,
                start_from_page=2,
                position="top_right",
                font_family="Times New Roman",
                font_size_pt=12.0,
            ),
        )
        return margins_dict, page_config

    def _extract_title_page_metadata(self, doc: Document) -> Tuple[Dict[str, Any], int]:
        """
        Ищет ключевые фразы титульного листа в первых параграфах документа.
        """
        meta: Dict[str, Any] = {}
        all_paras = doc.paragraphs[:40]  # Ограничиваем поиск первыми 40 абзацами
        
        has_title_page = False
        title_end_index = 0

        # Ключевые слова для детекции титульника
        title_triggers = [
            "МИНИСТЕРСТВО", "ДЕПАРТАМЕНТ", "УНИВЕРСИТЕТ", "КОЛЛЕДЖ", "ИНСТИТУТ",
            "КУРСОВОЙ ПРОЕКТ", "КУРСОВАЯ РАБОТА", "ДИПЛОМНЫЙ ПРОЕКТ", "РЕФЕРАТ", "ОТЧЕТ",
            "Выполнил", "Студент", "Руководитель", "Преподаватель"
        ]

        full_text_start = " ".join(p.text for p in all_paras)
        matches = [kw for kw in title_triggers if kw.lower() in full_text_start.lower()]
        
        if len(matches) >= 3:
            has_title_page = True

        if not has_title_page:
            return {}, 0

        # Поиск полей титульного листа регулярными выражениями и эвристиками
        for i, p in enumerate(all_paras):
            text = p.text.strip()
            if not text:
                continue

            # Проверяем на окончание титульника (город, год или разрыв страницы)
            if re.search(r'(Самара|Москва|Санкт-Петербург|Тольятти|Казань)\s*\n*.*(202\d|201\d)', text, re.IGNORECASE) or (len(text) < 20 and re.match(r'^202\d$', text)):
                title_end_index = i + 1
                year_match = re.search(r'202\d', text)
                if year_match:
                    meta["year"] = int(year_match.group(0))
                city_match = re.search(r'(Самара|Москва|Санкт-Петербург|Тольятти|Казань|Уфа|Саратов)', text, re.IGNORECASE)
                if city_match:
                    meta["city"] = city_match.group(0)

            # Министерство
            if "МИНИСТЕРСТВО" in text.upper():
                meta["ministry"] = text

            # Учебное заведение
            elif any(w in text.upper() for w in ["КОЛЛЕДЖ", "УНИВЕРСИТЕТ", "ИНСТИТУТ", "ТЕХНИКУМ", "ГБПОУ", "ФГБОУ"]):
                if "institution" not in meta:
                    meta["institution"] = text

            # Отделение / Кафедра / Факультет
            elif any(w in text.lower() for w in ["отделение", "кафедра", "факультет"]):
                if "кафедра" in text.lower():
                    meta["department"] = text
                elif "факультет" in text.lower():
                    meta["faculty"] = text
                else:
                    meta["department"] = text

            # Специальность
            elif re.search(r'\d{2}\.\d{2}\.\d{2}', text):
                meta["specialty"] = text

            # Тип работы
            elif any(w in text.upper() for w in ["КУРСОВОЙ ПРОЕКТ", "КУРСОВАЯ РАБОТА", "ДИПЛОМНЫЙ ПРОЕКТ", "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", "ОТЧЕТ ПО ПРАКТИКЕ", "РЕФЕРАТ"]):
                meta["work_type"] = text

            # Дисциплина / МДК
            elif any(w in text.upper() for w in ["МДК", "ДИСЦИПЛИНА", "ПО ДИСЦИПЛИНЕ", "ПО МДК"]):
                clean_disc = re.sub(r'^(по дисциплине|по мдк|дисциплина|мдк)[:\s]*', '', text, flags=re.IGNORECASE).strip()
                meta["discipline"] = clean_disc

            # Тема работы
            elif any(w in text.lower() for w in ["тема:", "на тему:"]):
                clean_theme = re.sub(r'^(тема|на тему)[:\s]*[«"]?', '', text, flags=re.IGNORECASE).rstrip('»"').strip()
                meta["theme"] = clean_theme

            # Студент
            elif any(w in text.lower() for w in ["выполнил", "студент", "обучающийся"]):
                student_match = re.search(r'([А-ЯЁ][а-яё]+)\s+([А-ЯЁ]\.\s*[А-ЯЁ]\.?)', text)
                if student_match:
                    meta["student_name"] = f"{student_match.group(1)} {student_match.group(2)}"
                group_match = re.search(r'гр\.?\s*([A-Za-zА-Яа-я0-9\-]+)', text)
                if group_match:
                    meta["student_group"] = f"гр. {group_match.group(1)}"

            # Руководитель
            elif any(w in text.lower() for w in ["руководитель", "преподаватель", "проверил"]):
                sup_match = re.search(r'([А-ЯЁ][а-яё]+)\s+([А-ЯЁ]\.\s*[А-ЯЁ]\.?)', text)
                if sup_match:
                    meta["supervisor_name"] = f"{sup_match.group(1)} {sup_match.group(2)}"
                if "преподаватель" in text.lower():
                    meta["supervisor_title"] = "преподаватель"
                elif "доцент" in text.lower():
                    meta["supervisor_title"] = "доцент"
                elif "профессор" in text.lower():
                    meta["supervisor_title"] = "профессор"

        # Также проверим таблицы титульного листа (сетка подписей)
        for t in doc.tables[:3]:
            for row in t.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if "студент" in cell_text.lower() or "выполнил" in cell_text.lower():
                        st_m = re.search(r'([А-ЯЁ][а-яё]+)\s+([А-ЯЁ]\.\s*[А-ЯЁ]\.?)', cell_text)
                        if st_m and "student_name" not in meta:
                            meta["student_name"] = f"{st_m.group(1)} {st_m.group(2)}"
                    if "руководитель" in cell_text.lower() or "проверил" in cell_text.lower():
                        sp_m = re.search(r'([А-ЯЁ][а-яё]+)\s+([А-ЯЁ]\.\s*[А-ЯЁ]\.?)', cell_text)
                        if sp_m and "supervisor_name" not in meta:
                            meta["supervisor_name"] = f"{sp_m.group(1)} {sp_m.group(2)}"

        if title_end_index == 0 and has_title_page:
            title_end_index = min(15, len(all_paras))

        return meta, title_end_index

    def _analyze_typography(self, doc: Document, skip_paras: int = 0) -> Tuple[TextStyle, HeadingStyles, TableStyle, ListStyle]:
        """
        Анализирует распределение шрифтов, кеглей и межстрочных интервалов в основном теле документа.
        """
        fonts_counter = Counter()
        sizes_counter = Counter()
        line_spacings_counter = Counter()
        first_indents_counter = Counter()
        alignments_counter = Counter()

        h1_style = TextStyle(
            font_family="Times New Roman", font_size_pt=14.0, bold=True, all_caps=True,
            alignment="CENTER", spacing_before_pt=12.0, spacing_after_pt=12.0, page_break_before=True, keep_with_next=True
        )
        h2_style = TextStyle(
            font_family="Times New Roman", font_size_pt=14.0, bold=True, all_caps=False,
            alignment="LEFT", first_line_indent_cm=1.25, spacing_before_pt=12.0, spacing_after_pt=6.0, keep_with_next=True
        )
        h3_style = TextStyle(
            font_family="Times New Roman", font_size_pt=14.0, bold=False, italic=True,
            alignment="LEFT", first_line_indent_cm=1.25, spacing_before_pt=6.0, spacing_after_pt=6.0, keep_with_next=True
        )

        body_paragraphs = doc.paragraphs[skip_paras:]
        for p in body_paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = (p.style.name or "").lower() if p.style else ""
            
            # Проверяем шрифты в Runs
            for r in p.runs:
                if r.text.strip():
                    if r.font.name:
                        fonts_counter[r.font.name] += len(r.text)
                    if r.font.size:
                        sizes_counter[round(r.font.size.pt, 1)] += len(r.text)

            # Выравнивание
            if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                alignments_counter["JUSTIFY"] += 1
            elif p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                alignments_counter["CENTER"] += 1
            elif p.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                alignments_counter["LEFT"] += 1
            elif p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                alignments_counter["RIGHT"] += 1

            # Межстрочный интервал
            if p.paragraph_format.line_spacing:
                try:
                    ls = float(p.paragraph_format.line_spacing)
                    line_spacings_counter[round(ls, 2)] += 1
                except Exception:
                    pass

            # Отступ первой строки
            if p.paragraph_format.first_line_indent:
                indent_cm = round(p.paragraph_format.first_line_indent.cm, 2)
                first_indents_counter[indent_cm] += 1

            # Проверяем H1 / H2 / H3
            r0 = p.runs[0] if p.runs else None
            is_bold = bool(r0 and r0.font.bold)
            run_font = r0.font.name if r0 and r0.font.name else None
            run_size = round(r0.font.size.pt, 1) if (r0 and r0.font.size) else None

            if "heading 1" in style_name or "заголовок 1" in style_name or (is_bold and p.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 70 and text.isupper()):
                h1_style.font_family = run_font or h1_style.font_family
                h1_style.font_size_pt = run_size or h1_style.font_size_pt
                h1_style.all_caps = text.isupper()
                h1_style.alignment = "CENTER" if p.alignment == WD_ALIGN_PARAGRAPH.CENTER else "LEFT"
                if p.paragraph_format.space_before:
                    h1_style.spacing_before_pt = round(p.paragraph_format.space_before.pt, 1)
                if p.paragraph_format.space_after:
                    h1_style.spacing_after_pt = round(p.paragraph_format.space_after.pt, 1)

            elif "heading 2" in style_name or "заголовок 2" in style_name or (is_bold and len(text) < 90 and re.match(r'^\d+\.\d+\s+', text)):
                h2_style.font_family = run_font or h2_style.font_family
                h2_style.font_size_pt = run_size or h2_style.font_size_pt
                h2_style.alignment = "LEFT"
                if p.paragraph_format.space_before:
                    h2_style.spacing_before_pt = round(p.paragraph_format.space_before.pt, 1)
                if p.paragraph_format.space_after:
                    h2_style.spacing_after_pt = round(p.paragraph_format.space_after.pt, 1)

        # Выбираем доминирующие параметры тела
        dominant_font = fonts_counter.most_common(1)[0][0] if fonts_counter else "Times New Roman"
        dominant_size = sizes_counter.most_common(1)[0][0] if sizes_counter else 14.0
        dominant_spacing = line_spacings_counter.most_common(1)[0][0] if line_spacings_counter else 1.5
        dominant_indent = first_indents_counter.most_common(1)[0][0] if first_indents_counter else 1.25
        dominant_align = alignments_counter.most_common(1)[0][0] if alignments_counter else "JUSTIFY"

        body_style = TextStyle(
            font_family=dominant_font,
            font_size_pt=dominant_size,
            line_spacing=dominant_spacing,
            first_line_indent_cm=dominant_indent,
            alignment=dominant_align, # type: ignore
            spacing_before_pt=0.0,
            spacing_after_pt=0.0,
        )

        heading_styles = HeadingStyles(h1=h1_style, h2=h2_style, h3=h3_style)

        table_style = TableStyle(
            font_family=dominant_font,
            font_size_pt=12.0 if dominant_size >= 14.0 else dominant_size,
            line_spacing=1.0,
            borders=True,
            repeat_header_on_new_page=True,
        )

        list_style = ListStyle(
            bullet_style="dash",
            first_line_indent_cm=dominant_indent,
            left_indent_cm=dominant_indent,
            spacing_before_pt=0.0,
            spacing_after_pt=0.0,
        )

        return body_style, heading_styles, table_style, list_style

    def _convert_document_to_markdown(
        self, doc: Document, title_meta: Dict[str, Any], profile: DocProfile, skip_paras: int
    ) -> Tuple[str, List[str], int]:
        """
        Транслирует параграфы и таблицы документа в семантический Markdown с YAML Frontmatter.
        """
        md_lines: List[str] = []
        headings_hierarchy: List[str] = []
        tables_count = len(doc.tables)

        # 1. Формируем YAML Frontmatter
        frontmatter_dict: Dict[str, Any] = {
            "profile": profile.meta.name,
        }
        if title_meta:
            frontmatter_dict["title"] = title_meta

        # Собираем оглавление
        toc_entries = []

        # 2. Итерируем элементы документа
        body_elements = doc.paragraphs[skip_paras:]
        in_code_block = False

        for p in body_elements:
            text = p.text.strip()
            if not text:
                continue

            style_name = (p.style.name or "").lower() if p.style else ""
            r0 = p.runs[0] if p.runs else None
            is_bold = bool(r0 and r0.font.bold)

            # Проверяем H1
            if "heading 1" in style_name or "заголовок 1" in style_name or (is_bold and p.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 70 and (text.isupper() or re.match(r'^\d+\s+[А-ЯЁA-Z]', text))):
                clean_title = text.strip()
                md_lines.append(f"\n# {clean_title}\n")
                headings_hierarchy.append(f"H1: {clean_title}")
                toc_entries.append({"title": clean_title, "level": 1})
                continue

            # Проверяем H2
            if "heading 2" in style_name or "заголовок 2" in style_name or (is_bold and len(text) < 90 and re.match(r'^\d+\.\d+\s+', text)):
                clean_title = text.strip()
                md_lines.append(f"\n## {clean_title}\n")
                headings_hierarchy.append(f"H2: {clean_title}")
                toc_entries.append({"title": clean_title, "level": 2})
                continue

            # Проверяем H3
            if "heading 3" in style_name or "заголовок 3" in style_name or (len(text) < 90 and re.match(r'^\d+\.\d+\.\d+\s+', text)):
                clean_title = text.strip()
                md_lines.append(f"\n### {clean_title}\n")
                headings_hierarchy.append(f"H3: {clean_title}")
                toc_entries.append({"title": clean_title, "level": 3})
                continue

            # Проверяем списки
            if "list" in style_name or text.startswith("- ") or text.startswith("• "):
                clean_item = re.sub(r'^[•\-\–—]\s*', '', text)
                md_lines.append(f"- {clean_item}")
                continue

            # Проверяем нумерованные списки
            if re.match(r'^\d+[\.\)]\s+', text):
                md_lines.append(text)
                continue

            # Форматируем Run-ы внутри абзаца (жирный, курсив, код)
            formatted_para = self._format_paragraph_runs(p)
            md_lines.append(formatted_para)

        # 3. Конвертируем таблицы
        for table_idx, table in enumerate(doc.tables):
            table_md = self._convert_table_to_gfm(table, table_idx + 1)
            md_lines.append(f"\n{table_md}\n")

        if toc_entries:
            frontmatter_dict["toc"] = toc_entries

        # Собираем итоговый Markdown
        fm_yaml = yaml.dump(frontmatter_dict, allow_unicode=True, sort_keys=False).strip()
        full_md = f"---\n{fm_yaml}\n---\n\n" + "\n\n".join(md_lines)
        # Очищаем множественные пустые строки
        full_md = re.sub(r'\n{3,}', '\n\n', full_md)
        return full_md, headings_hierarchy, tables_count

    def _format_paragraph_runs(self, p) -> str:
        """
        Преобразует Run-ы docx с сохранением жирного, курсива и моноширинного шрифта.
        """
        parts: List[str] = []
        for r in p.runs:
            t = r.text
            if not t:
                continue
            
            # Сохраняем пробелы по краям
            l_spaces = len(t) - len(t.lstrip())
            r_spaces = len(t) - len(t.rstrip())
            core = t.strip()
            if not core:
                parts.append(t)
                continue

            # Оформление
            if r.font.name in ["Courier New", "Consolas", "Courier"]:
                core = f"`{core}`"
            elif r.font.bold and r.font.italic:
                core = f"***{core}***"
            elif r.font.bold:
                core = f"**{core}**"
            elif r.font.italic:
                core = f"*{core}*"

            parts.append((" " * l_spaces) + core + (" " * r_spaces))

        return "".join(parts) if parts else p.text

    def _convert_table_to_gfm(self, table, table_idx: int) -> str:
        """
        Преобразует таблицу DOCX в Markdown таблицу GitHub Flavored Markdown (GFM).
        """
        rows = table.rows
        if not rows:
            return ""

        table_lines: List[str] = []
        
        # Получаем данные ячеек
        grid: List[List[str]] = []
        for r in rows:
            row_data = [c.text.strip().replace("\n", " ") for c in r.cells]
            grid.append(row_data)

        if not grid:
            return ""

        # Шапка таблицы
        headers = grid[0]
        num_cols = len(headers)
        table_lines.append("| " + " | ".join(headers) + " |")
        table_lines.append("| " + " | ".join([":---"] * num_cols) + " |")

        # Тело таблицы
        for row in grid[1:]:
            # Выравниваем количество колонок при объединениях
            padded_row = row + [""] * (num_cols - len(row))
            table_lines.append("| " + " | ".join(padded_row[:num_cols]) + " |")

        return "\n".join(table_lines)

    def _generate_layout_report(
        self, profile: DocProfile, title_meta: Dict[str, Any], tables_count: int, headings_count: int
    ) -> str:
        """
        Генерирует человеко- и машино-читаемый отчет о точных параметрах макета документа.
        """
        p = profile
        m = p.page.margins
        lines = [
            f"📐 ГЕОМЕТРИЯ СТРАНИЦЫ: Формат {p.page.format} ({p.page.orientation})",
            f"   • Поля: Левое {m.left_mm} мм (под переплет), Правое {m.right_mm} мм, Верхнее {m.top_mm} мм, Нижнее {m.bottom_mm} мм",
            f"   • Нумерация: {'Включена' if p.page.page_numbering.enabled else 'Выключена'} (с {p.page.page_numbering.start_from_page} стр., {p.page.page_numbering.position})",
            "",
            f"✍️ ТИПОГРАФИКА ОСНОВНОГО ТЕКСТА:",
            f"   • Гарнитура: {p.body.font_family}, Кегль: {p.body.font_size_pt} pt",
            f"   • Межстрочный интервал: {p.body.line_spacing} строки",
            f"   • Абзацный отступ (красная строка): {p.body.first_line_indent_cm} см",
            f"   • Выравнивание: {p.body.alignment}",
            f"   • Интервалы абзаца: До {p.body.spacing_before_pt} pt, После {p.body.spacing_after_pt} pt",
            "",
            f"🏷️ ИЕРАРХИЯ ЗАГОЛОВКОВ (Найдено {headings_count} шт.):",
            f"   • H1: {p.headings.h1.font_family} {p.headings.h1.font_size_pt} pt, Bold={'Да' if p.headings.h1.bold else 'Нет'}, Caps={'Да' if p.headings.h1.all_caps else 'Нет'}, Align={p.headings.h1.alignment}, Отступы До={p.headings.h1.spacing_before_pt}pt / После={p.headings.h1.spacing_after_pt}pt",
            f"   • H2: {p.headings.h2.font_family} {p.headings.h2.font_size_pt} pt, Bold={'Да' if p.headings.h2.bold else 'Нет'}, Align={p.headings.h2.alignment}, Отступы До={p.headings.h2.spacing_before_pt}pt / После={p.headings.h2.spacing_after_pt}pt",
            f"   • H3: {p.headings.h3.font_family} {p.headings.h3.font_size_pt} pt, Italic={'Да' if p.headings.h3.italic else 'Нет'}, Align={p.headings.h3.alignment}",
            "",
            f"📊 ТАБЛИЦЫ И СПИСКИ:",
            f"   • Таблиц в документе: {tables_count} шт. (Шрифт: {p.elements.tables.font_family} {p.elements.tables.font_size_pt} pt, Интервал: {p.elements.tables.line_spacing})",
            f"   • Списки: Маркер '{p.elements.lists.bullet_style}', Красная строка: {p.elements.lists.first_line_indent_cm} см",
        ]

        if title_meta:
            lines.append("")
            lines.append("🎓 ИЗВЛЕЧЕННЫЙ ТИТУЛЬНЫЙ ЛИСТ:")
            for k, v in title_meta.items():
                lines.append(f"   • {k}: {v}")

        return "\n".join(lines)
